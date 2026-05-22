
from typing import List
import math
from copy import deepcopy
from pathlib import Path
from tempfile import TemporaryDirectory

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .config import ATTAINMENT_THRESHOLD

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
except Exception:
    plt = None


def export_result_workbook(output_path: Path, student_target_df: pd.DataFrame, course_target_df: pd.DataFrame) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        student_target_df.to_excel(writer, sheet_name="学生课程目标结果表", index=False)
        course_target_df.to_excel(writer, sheet_name="课程目标结果表", index=False)


ANALYSIS_PASS_THRESHOLD = ATTAINMENT_THRESHOLD


def _build_course_summary_sheet(course_target_df: pd.DataFrame) -> pd.DataFrame:
    df = course_target_df.copy()
    result = df[
        [
            "课程目标编号",
            "课程目标描述",
            "综合平均达成值",
        ]
    ].copy()
    result["合格阈值"] = ANALYSIS_PASS_THRESHOLD
    result["是否达成"] = result["综合平均达成值"].apply(
        lambda v: "达成" if pd.notna(v) and v >= ANALYSIS_PASS_THRESHOLD else "未达成"
    )
    return result


def _build_student_attainment_sheet(student_target_df: pd.DataFrame) -> pd.DataFrame:
    df = student_target_df.copy()
    result = df[
        [
            "学号",
            "姓名",
            "课程目标编号",
            "过程性贡献值",
            "期末贡献值",
            "综合达成值",
        ]
    ].copy()
    result = result.rename(
        columns={
            "过程性贡献值": "过程性成绩",
            "期末贡献值": "期末成绩",
            "综合达成值": "综合达成度",
        }
    )
    result["是否达成"] = result["综合达成度"].apply(
        lambda v: "达成" if pd.notna(v) and v >= ANALYSIS_PASS_THRESHOLD else "未达成"
    )
    return result


def _build_unmet_students_sheet(student_target_df: pd.DataFrame) -> pd.DataFrame:
    student_attainment = _build_student_attainment_sheet(student_target_df)
    result = student_attainment[student_attainment["综合达成度"] < ANALYSIS_PASS_THRESHOLD].copy()
    result = result[
        [
            "课程目标编号",
            "学号",
            "姓名",
            "综合达成度",
        ]
    ]
    return result.sort_values(["课程目标编号", "学号"]).reset_index(drop=True)


def _attainment_interval(v) -> str:
    if pd.isna(v):
        return ""
    if v < ANALYSIS_PASS_THRESHOLD:
        return f"<{ANALYSIS_PASS_THRESHOLD:.2f}"
    if v < 0.70:
        return f"{ANALYSIS_PASS_THRESHOLD:.2f}-0.70"
    if v < 0.80:
        return "0.70-0.80"
    if v < 0.90:
        return "0.80-0.90"
    return ">=0.90"


def _build_distribution_sheet(student_target_df: pd.DataFrame) -> pd.DataFrame:
    df = student_target_df[["课程目标编号", "综合达成值"]].copy()
    df["达成度区间"] = df["综合达成值"].apply(_attainment_interval)
    intervals = [f"<{ANALYSIS_PASS_THRESHOLD:.2f}", f"{ANALYSIS_PASS_THRESHOLD:.2f}-0.70", "0.70-0.80", "0.80-0.90", ">=0.90"]
    rows = []

    for target_id, grp in df.groupby("课程目标编号", sort=True):
        total = len(grp)
        counts = grp["达成度区间"].value_counts().to_dict()
        for interval in intervals:
            count = int(counts.get(interval, 0))
            rows.append(
                {
                    "课程目标编号": target_id,
                    "达成度区间": interval,
                    "人数": count,
                    "占比": count / total if total else 0,
                }
            )

    return pd.DataFrame(rows)


def _build_chart_data_sheet(course_target_df: pd.DataFrame, distribution_df: pd.DataFrame) -> pd.DataFrame:
    avg_df = course_target_df[["课程目标编号", "综合平均达成值"]].copy()
    unmet_df = distribution_df[distribution_df["达成度区间"] == f"<{ANALYSIS_PASS_THRESHOLD:.2f}"][
        ["课程目标编号", "人数"]
    ].rename(columns={"人数": "未达成人数"})
    interval_counts = distribution_df.pivot_table(
        index="课程目标编号",
        columns="达成度区间",
        values="人数",
        aggfunc="sum",
        fill_value=0,
    ).reset_index()

    result = avg_df.merge(unmet_df, on="课程目标编号", how="left")
    result = result.merge(interval_counts, on="课程目标编号", how="left")
    result["未达成人数"] = result["未达成人数"].fillna(0).astype(int)
    return result


def export_analysis_workbook(
    output_path: Path,
    analysis_df: pd.DataFrame,
    student_target_df: pd.DataFrame = None,
    course_target_df: pd.DataFrame = None,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        analysis_df.to_excel(writer, sheet_name="分析底表", index=False)
        if student_target_df is not None and course_target_df is not None:
            course_summary_df = _build_course_summary_sheet(course_target_df)
            student_attainment_df = _build_student_attainment_sheet(student_target_df)
            unmet_students_df = _build_unmet_students_sheet(student_target_df)
            distribution_df = _build_distribution_sheet(student_target_df)
            chart_data_df = _build_chart_data_sheet(course_target_df, distribution_df)

            course_summary_df.to_excel(writer, sheet_name="课程目标汇总表", index=False)
            student_attainment_df.to_excel(writer, sheet_name="学生个体达成度表", index=False)
            unmet_students_df.to_excel(writer, sheet_name="未达成学生清单", index=False)
            distribution_df.to_excel(writer, sheet_name="达成度分布统计表", index=False)
            chart_data_df.to_excel(writer, sheet_name="图表数据表", index=False)


def _normalize_text(v) -> str:
    if pd.isna(v):
        return ""
    s = str(v).strip().replace("\u3000", " ")
    if s.endswith(".0"):
        s = s[:-2]
    return s


def _normalize_target_id(v) -> str:
    s = _normalize_text(v)
    if not s:
        return ""
    if s.startswith("课程目标"):
        return s
    if s.isdigit():
        return f"课程目标{s}"
    return s


def _fmt_score(v, digits=1) -> str:
    if pd.isna(v):
        return ""
    return f"{float(v):.{digits}f}"


def _fmt_attainment(v, digits=2) -> str:
    if pd.isna(v):
        return ""
    return f"{float(v):.{digits}f}"


def _find_template_docx(course_path: Path) -> Path:
    candidates = [p for p in course_path.glob("*.docx") if not p.name.startswith("~$")]
    if not candidates:
        raise FileNotFoundError(f"在 {course_path} 下未找到报告模板 docx 文件。")
    preferred = [p for p in candidates if "达成情况分析" in p.stem or "持续改进报告" in p.stem]
    return sorted(preferred or candidates)[0]


def _clear_cell_keep_style(cell, text: str, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if not cell.paragraphs:
        cell.text = text
        return
    # 保留第一个段落的样式
    p = cell.paragraphs[0]
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)
    runs = list(p.runs)
    if runs:
        # 清空首个 run，删除其余 run
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
        runs[0].text = text
    else:
        p.add_run(text)
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _replace_data_rows_with_clones(table, start_row_idx: int, row_values: List[List[str]]) -> None:
    if len(table.rows) <= start_row_idx:
        raise ValueError("模板表格中没有可供复制的数据行。")
    prototype_tr = deepcopy(table.rows[start_row_idx]._tr)

    # 删除原有数据行
    while len(table.rows) > start_row_idx:
        table._tbl.remove(table.rows[start_row_idx]._tr)

    # 逐行克隆写入
    for values in row_values:
        new_tr = deepcopy(prototype_tr)
        table._tbl.append(new_tr)
        row = table.rows[-1]
        for i, value in enumerate(values):
            if i < len(row.cells):
                align = WD_ALIGN_PARAGRAPH.CENTER
                if i == 2:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                _clear_cell_keep_style(row.cells[i], "" if value is None else str(value), align=align)


def _build_table1_wide(student_target_df: pd.DataFrame) -> pd.DataFrame:
    df = student_target_df.copy()
    df["课程目标编号"] = df["课程目标编号"].apply(_normalize_target_id)
    df["学号"] = df["学号"].apply(_normalize_text)
    df["姓名"] = df["姓名"].apply(_normalize_text)

    needed = ["学号", "姓名", "课程目标编号", "平时作业", "实验", "实验报告", "课堂表现", "试卷得分", "综合达成值"]
    missing = [c for c in needed if c not in df.columns]
    if missing:
        raise ValueError(f"student_target_df 缺少字段: {missing}")

    wide = df[needed].pivot_table(
        index=["学号", "姓名"],
        columns="课程目标编号",
        values=["平时作业", "实验", "实验报告", "课堂表现", "试卷得分", "综合达成值"],
        aggfunc="first",
    )
    wide.columns = [f"{target}_{metric}" for metric, target in wide.columns]
    wide = wide.reset_index().sort_values(["学号", "姓名"]).reset_index(drop=True)
    wide.insert(0, "序号", range(1, len(wide) + 1))
    return wide


def _fill_table1(table, student_target_df: pd.DataFrame) -> None:
    # 模板首表当前实际为 15 列：
    # 序号 学号 姓名 + 课程目标1(平时/实验/实验报告/课堂/试卷/达成度) + 课程目标2(...)
    if len(table.columns) != 15:
        raise ValueError(f"模板中的第一个表列数为 {len(table.columns)}，当前脚本按 15 列模板处理。")

    wide = _build_table1_wide(student_target_df)

    def gv(row, col):
        return row[col] if col in row.index else None

    values = []
    for _, row in wide.iterrows():
        values.append([
            _normalize_text(row["序号"]),
            _normalize_text(row["学号"]),
            _normalize_text(row["姓名"]),
            _fmt_score(gv(row, "课程目标1_平时作业"), 1),
            _fmt_score(gv(row, "课程目标1_实验"), 1),
            _fmt_score(gv(row, "课程目标1_实验报告"), 1),
            _fmt_score(gv(row, "课程目标1_课堂表现"), 1),
            _fmt_score(gv(row, "课程目标1_试卷得分"), 1),
            _fmt_attainment(gv(row, "课程目标1_综合达成值"), 2),
            _fmt_score(gv(row, "课程目标2_平时作业"), 1),
            _fmt_score(gv(row, "课程目标2_实验"), 1),
            _fmt_score(gv(row, "课程目标2_实验报告"), 1),
            _fmt_score(gv(row, "课程目标2_课堂表现"), 1),
            _fmt_score(gv(row, "课程目标2_试卷得分"), 1),
            _fmt_attainment(gv(row, "课程目标2_综合达成值"), 2),
        ])

    avg_row = [
        "", "", "平均值",
        _fmt_score(wide["课程目标1_平时作业"].mean() if "课程目标1_平时作业" in wide.columns else None, 1),
        _fmt_score(wide["课程目标1_实验"].mean() if "课程目标1_实验" in wide.columns else None, 1),
        _fmt_score(wide["课程目标1_实验报告"].mean() if "课程目标1_实验报告" in wide.columns else None, 1),
        _fmt_score(wide["课程目标1_课堂表现"].mean() if "课程目标1_课堂表现" in wide.columns else None, 1),
        _fmt_score(wide["课程目标1_试卷得分"].mean() if "课程目标1_试卷得分" in wide.columns else None, 1),
        _fmt_attainment(wide["课程目标1_综合达成值"].mean() if "课程目标1_综合达成值" in wide.columns else None, 2),
        _fmt_score(wide["课程目标2_平时作业"].mean() if "课程目标2_平时作业" in wide.columns else None, 1),
        _fmt_score(wide["课程目标2_实验"].mean() if "课程目标2_实验" in wide.columns else None, 1),
        _fmt_score(wide["课程目标2_实验报告"].mean() if "课程目标2_实验报告" in wide.columns else None, 1),
        _fmt_score(wide["课程目标2_课堂表现"].mean() if "课程目标2_课堂表现" in wide.columns else None, 1),
        _fmt_score(wide["课程目标2_试卷得分"].mean() if "课程目标2_试卷得分" in wide.columns else None, 1),
        _fmt_attainment(wide["课程目标2_综合达成值"].mean() if "课程目标2_综合达成值" in wide.columns else None, 2),
    ]
    values.append(avg_row)

    _replace_data_rows_with_clones(table, start_row_idx=3, row_values=values)


def _fill_table2(table, analysis_df: pd.DataFrame) -> None:
    if len(table.rows) < 4 or len(table.columns) < 3:
        raise ValueError("模板中的第二个表结构不符合预期。")
    temp = analysis_df.copy()
    temp["课程目标编号"] = temp["课程目标编号"].apply(_normalize_target_id)
    m = {r["课程目标编号"]: r for _, r in temp.iterrows()}
    t1 = m.get("课程目标1")
    t2 = m.get("课程目标2")
    _clear_cell_keep_style(table.cell(3, 1), _fmt_attainment(t1["综合平均达成值"] if t1 is not None else None, 2))
    _clear_cell_keep_style(table.cell(3, 2), _fmt_attainment(t2["综合平均达成值"] if t2 is not None else None, 2))


def _read_table2_series(table) -> dict:
    grades, y1, y2 = [], [], []
    for r in range(1, len(table.rows)):
        g = _normalize_text(table.cell(r, 0).text)
        if not g:
            continue
        grades.append(g)
        try:
            y1.append(float(_normalize_text(table.cell(r, 1).text)))
        except Exception:
            y1.append(float("nan"))
        try:
            y2.append(float(_normalize_text(table.cell(r, 2).text)))
        except Exception:
            y2.append(float("nan"))
    return {"grades": grades, "课程目标1": y1, "课程目标2": y2}


def _prepare_matplotlib():
    if plt is None:
        raise ImportError("当前环境缺少 matplotlib，请先执行：python -m pip install matplotlib")
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def _plot_three_year_compare(table2, save_path: Path) -> None:
    _prepare_matplotlib()
    s = _read_table2_series(table2)
    x = list(range(len(s["grades"])))
    width = 0.35
    fig, ax = plt.subplots(figsize=(8.0, 4.6))
    ax.bar([i - width / 2 for i in x], s["课程目标1"], width=width, label="课程目标1")
    ax.bar([i + width / 2 for i in x], s["课程目标2"], width=width, label="课程目标2")
    ax.axhline(ATTAINMENT_THRESHOLD, linestyle="--", linewidth=1, label="达标阈值")
    ax.set_xticks(x)
    ax.set_xticklabels(s["grades"])
    ax.set_ylabel("达成值")
    ax.legend()
    fig.tight_layout()
    fig.savefig(save_path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def _plot_target_scatter(student_target_df: pd.DataFrame, target_id: str, save_path: Path) -> None:
    _prepare_matplotlib()
    df = student_target_df.copy()
    df["课程目标编号"] = df["课程目标编号"].apply(_normalize_target_id)
    df = df[df["课程目标编号"] == target_id].copy().sort_values(["学号", "姓名"]).reset_index(drop=True)
    y = pd.to_numeric(df["综合达成值"], errors="coerce")
    x = list(range(1, len(df) + 1))

    fig, ax = plt.subplots(figsize=(8.0, 4.6))
    ax.scatter(x, y)
    ax.axhline(ATTAINMENT_THRESHOLD, linestyle="--", linewidth=1, label="达标阈值")
    ax.set_xlabel("学生序号")
    ax.set_ylabel("综合达成值")
    ax.legend()
    fig.tight_layout()
    fig.savefig(save_path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def _insert_picture_after_paragraph(paragraph, img_path: Path, width_cm: float = 15.5):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    pic_para = paragraph._parent.add_paragraph()
    pic_para._p.getparent().remove(pic_para._p)
    new_p.getparent().replace(new_p, pic_para._p)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    pic_para.paragraph_format.space_before = Pt(4)
    pic_para.paragraph_format.space_after = Pt(10)


def _find_paragraph_contains(doc: Document, text: str):
    for p in doc.paragraphs:
        if text in p.text:
            return p
    return None


def export_report_docx(
    output_path: Path,
    course_df: pd.DataFrame,
    analysis_df: pd.DataFrame,
    student_target_df: pd.DataFrame,
    course_path: Path,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    template_path = _find_template_docx(course_path)
    doc = Document(template_path)

    if len(doc.tables) < 2:
        raise ValueError("模板文档中至少需要包含2个表。")

    _fill_table1(doc.tables[0], student_target_df)
    _fill_table2(doc.tables[1], analysis_df)

    with TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        fig1 = tmpdir / "fig1.png"
        fig2 = tmpdir / "fig2.png"
        fig3 = tmpdir / "fig3.png"

        _plot_three_year_compare(doc.tables[1], fig1)
        _plot_target_scatter(student_target_df, "课程目标1", fig2)
        _plot_target_scatter(student_target_df, "课程目标2", fig3)

        # 将图插到对应图题后面。模板原图片保留时可能重复，因此先不尝试复杂替换，直接插入在图题后。
        # 如果模板中已有旧图，建议先手工删除旧图，仅保留图题。
        p1 = _find_paragraph_contains(doc, "图1")
        p2 = _find_paragraph_contains(doc, "图2")
        p3 = _find_paragraph_contains(doc, "图3")

        if p3 is not None:
            _insert_picture_after_paragraph(p3, fig3, width_cm=15.5)
        if p2 is not None:
            _insert_picture_after_paragraph(p2, fig2, width_cm=15.5)
        if p1 is not None:
            _insert_picture_after_paragraph(p1, fig1, width_cm=15.5)
        # 倒序插入，避免后插入时段落引用位移

        doc.save(output_path)
