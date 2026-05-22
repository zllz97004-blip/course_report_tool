from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPORT_PASS_THRESHOLD = 0.60


def _fmt(v, digits=6) -> str:
    if pd.isna(v):
        return ""
    return f"{float(v):.{digits}f}"


def _set_document_font(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2"]:
        if style_name in styles:
            styles[style_name].font.name = "宋体"
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_table(doc: Document, df: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            value = row[col]
            if isinstance(value, float):
                cells[i].text = _fmt(value)
            else:
                cells[i].text = "" if pd.isna(value) else str(value)


def _build_course_summary(course_target_df: pd.DataFrame) -> pd.DataFrame:
    df = course_target_df[
        [
            "课程目标编号",
            "课程目标描述",
            "综合平均达成值",
        ]
    ].copy()
    df["合格阈值"] = REPORT_PASS_THRESHOLD
    df["是否达成"] = df["综合平均达成值"].apply(
        lambda v: "达成" if pd.notna(v) and v >= REPORT_PASS_THRESHOLD else "未达成"
    )
    return df


def _build_student_attainment(student_target_df: pd.DataFrame) -> pd.DataFrame:
    df = student_target_df[
        [
            "学号",
            "姓名",
            "课程目标编号",
            "过程性贡献值",
            "期末贡献值",
            "综合达成值",
        ]
    ].copy()
    df = df.rename(
        columns={
            "过程性贡献值": "过程性成绩",
            "期末贡献值": "期末成绩",
            "综合达成值": "综合达成度",
        }
    )
    df["是否达成"] = df["综合达成度"].apply(
        lambda v: "达成" if pd.notna(v) and v >= REPORT_PASS_THRESHOLD else "未达成"
    )
    return df


def _build_unmet_students(student_attainment_df: pd.DataFrame) -> pd.DataFrame:
    return student_attainment_df[
        student_attainment_df["综合达成度"] < REPORT_PASS_THRESHOLD
    ].copy()


def _build_distribution(student_attainment_df: pd.DataFrame) -> pd.DataFrame:
    def interval(v):
        if pd.isna(v):
            return ""
        if v < 0.60:
            return "<0.60"
        if v < 0.70:
            return "0.60-0.70"
        if v < 0.80:
            return "0.70-0.80"
        if v < 0.90:
            return "0.80-0.90"
        return ">=0.90"

    intervals = ["<0.60", "0.60-0.70", "0.70-0.80", "0.80-0.90", ">=0.90"]
    df = student_attainment_df[["课程目标编号", "综合达成度"]].copy()
    df["达成度区间"] = df["综合达成度"].apply(interval)

    rows = []
    for target_id, grp in df.groupby("课程目标编号", sort=True):
        total = len(grp)
        counts = grp["达成度区间"].value_counts().to_dict()
        for item in intervals:
            count = int(counts.get(item, 0))
            rows.append(
                {
                    "课程目标编号": target_id,
                    "达成度区间": item,
                    "人数": count,
                    "占比": count / total if total else 0,
                }
            )
    return pd.DataFrame(rows)


def _add_basic_info(doc: Document, course_df: pd.DataFrame) -> None:
    _add_heading(doc, "一、课程基本信息")
    row = course_df.iloc[0]
    info = pd.DataFrame(
        [
            {"项目": "课程名称", "内容": row.get("课程名称", "")},
            {"项目": "课程代码", "内容": row.get("课程代码", "")},
            {"项目": "开课学期", "内容": row.get("开课学期", "")},
        ]
    )
    _add_table(doc, info)


def _add_course_targets(doc: Document, course_df: pd.DataFrame) -> None:
    _add_heading(doc, "二、课程目标及支撑关系")
    df = course_df[["课程目标编号", "课程目标描述"]].copy()
    df["毕业要求支撑关系"] = "待完善"
    _add_table(doc, df)


def _add_method(doc: Document) -> None:
    _add_heading(doc, "三、课程目标达成度评价方法")
    doc.add_paragraph(
        "本课程采用过程性考核与期末考试相结合的方式评价课程目标达成情况。"
        "过程性考核成绩 = 课堂表现×10% + 课程作业×50% + 实验操作×20% + 实验报告×20%。"
        "个体达成度合格阈值为 0.60。"
    )


def _add_result(doc: Document, course_summary_df: pd.DataFrame) -> None:
    _add_heading(doc, "四、课程目标达成度结果")
    _add_table(doc, course_summary_df)


def _add_charts(doc: Document, chart_paths: dict = None) -> None:
    if not chart_paths:
        return

    chart_items = [
        ("图1 3年课程目标达成度对比图", chart_paths.get("three_year_compare")),
    ]
    for idx, (target_id, path) in enumerate(chart_paths.get("scatter", []), start=2):
        chart_items.append((f"图{idx} {target_id}达成度值散点图", path))

    _add_heading(doc, "五、持续改进报告图表")
    for caption, path in chart_items:
        if path is None or not Path(path).exists():
            continue
        doc.add_picture(str(path), width=Cm(15))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph = doc.add_paragraph(caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_student_analysis(
    doc: Document,
    course_summary_df: pd.DataFrame,
    student_attainment_df: pd.DataFrame,
    unmet_students_df: pd.DataFrame,
) -> None:
    _add_heading(doc, "六、学生个体达成情况分析")
    for _, row in course_summary_df.iterrows():
        target_id = row["课程目标编号"]
        total = int((student_attainment_df["课程目标编号"] == target_id).sum())
        unmet_count = int((unmet_students_df["课程目标编号"] == target_id).sum())
        unmet_text = (
            "但仍有部分学生未达到个体达成要求，说明学生在相关知识点掌握和综合应用方面仍存在差异。"
            if unmet_count > 0
            else "学生个体达成情况整体较为稳定，后续仍需保持过程性跟踪。"
        )
        doc.add_paragraph(
            f"{target_id}平均达成值为{_fmt(row['综合平均达成值'])}，"
            f"{'达到' if row['是否达成'] == '达成' else '未达到'}0.60的合格阈值。"
            f"参与评价学生记录数为{total}，未达成学生记录数为{unmet_count}。"
            f"{unmet_text}"
        )


def _add_problem_analysis(
    doc: Document,
    course_summary_df: pd.DataFrame,
    distribution_df: pd.DataFrame,
    unmet_students_df: pd.DataFrame,
) -> None:
    _add_heading(doc, "七、存在问题")
    for _, row in course_summary_df.iterrows():
        target_id = row["课程目标编号"]
        target_dist = distribution_df[distribution_df["课程目标编号"] == target_id]
        low_count = int(target_dist.loc[target_dist["达成度区间"] == "<0.60", "人数"].sum())
        near_count = int(target_dist.loc[target_dist["达成度区间"] == "0.60-0.70", "人数"].sum())
        unmet_count = int((unmet_students_df["课程目标编号"] == target_id).sum())
        doc.add_paragraph(
            f"{target_id}整体评价结果为{row['是否达成']}，但从学生个体分布看，"
            f"低于0.60的学生记录数为{low_count}，处于0.60-0.70区间的学生记录数为{near_count}。"
            f"这表明部分学生对相关知识点的理解深度、工程问题分析能力或综合应用能力仍有提升空间，"
            f"后续教学中应结合未达成学生清单开展更有针对性的过程反馈与学习支持。"
        )


def _add_improvement(doc: Document) -> None:
    _add_heading(doc, "八、持续改进措施")
    measures = [
        "加强薄弱知识点讲解，结合课堂练习和阶段性检测及时发现学生理解偏差。",
        "优化过程性考核反馈机制，对作业、实验操作和实验报告中的共性问题进行集中讲评。",
        "强化典型工程案例训练，并对未达成学生开展针对性辅导，提升知识迁移和综合应用能力。",
    ]
    for item in measures:
        doc.add_paragraph(item, style="List Number")


def _add_conclusion(doc: Document, course_summary_df: pd.DataFrame) -> None:
    _add_heading(doc, "九、结论")
    reached = course_summary_df[course_summary_df["是否达成"] == "达成"]["课程目标编号"].tolist()
    unreached = course_summary_df[course_summary_df["是否达成"] != "达成"]["课程目标编号"].tolist()
    parts = []
    if reached:
        parts.append(f"{'、'.join(reached)}达到0.60的评价标准")
    if unreached:
        parts.append(f"{'、'.join(unreached)}暂未达到0.60的评价标准")
    doc.add_paragraph(
        f"综合评价结果显示，{'; '.join(parts)}。"
        "后续课程教学将继续围绕课程目标达成情况，持续改进教学组织、过程评价和学习支持方式。"
    )


def export_report_docx(
    output_path: Path,
    course_df: pd.DataFrame,
    course_target_df: pd.DataFrame,
    student_target_df: pd.DataFrame,
    chart_paths: dict = None,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_document_font(doc)

    title = doc.add_heading("课程目标达成度报告草稿", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    course_summary_df = _build_course_summary(course_target_df)
    student_attainment_df = _build_student_attainment(student_target_df)
    unmet_students_df = _build_unmet_students(student_attainment_df)
    distribution_df = _build_distribution(student_attainment_df)

    _add_basic_info(doc, course_df)
    _add_course_targets(doc, course_df)
    _add_method(doc)
    _add_result(doc, course_summary_df)
    _add_charts(doc, chart_paths)
    _add_student_analysis(doc, course_summary_df, student_attainment_df, unmet_students_df)
    _add_problem_analysis(doc, course_summary_df, distribution_df, unmet_students_df)
    _add_improvement(doc)
    _add_conclusion(doc, course_summary_df)

    doc.save(output_path)
